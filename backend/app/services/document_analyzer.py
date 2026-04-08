"""
Document Analyzer Service
Analyzes uploaded documents and recommends optimal RAG configurations.
Supports PDFs, plain text, code, markdown, DOCX, and more.
"""
import re
from pathlib import Path
from typing import Dict, List, Optional
import numpy as np
from transformers import pipeline

from app.core.logging import get_logger
from app.services.pdf_processor import pdf_processor

logger = get_logger(__name__)

# Extensions that can be read as plain UTF-8 text
_TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".py", ".js", ".ts", ".tsx", ".jsx",
    ".json", ".csv", ".html", ".htm", ".xml", ".yaml", ".yml",
    ".java", ".cpp", ".c", ".h", ".go", ".rs", ".rb", ".php",
    ".swift", ".kt", ".scala", ".r", ".sql", ".sh", ".bash", ".zsh",
    ".ps1", ".css", ".scss", ".less", ".toml", ".ini", ".cfg",
    ".env", ".properties", ".rst", ".tex", ".rtf", ".log",
    ".code", ".config",
}


class DocumentAnalyzer:
    """
    Intelligent document analysis service that classifies documents,
    analyzes structure and density, and recommends RAG configurations.
    """
    
    def __init__(self):
        """Initialize the analyzer with ML models."""
        self._classifier = None  # Lazy load
        self._initialized = False
        
    def _ensure_initialized(self):
        """Lazy load the classification model to avoid startup delays."""
        if not self._initialized:
            logger.info("Loading fast zero-shot classification model (distilbart)...")
            try:
                # Use a much smaller and faster model
                self._classifier = pipeline(
                    "zero-shot-classification",
                    model="valhalla/distilbart-mnli-12-1",
                    device=-1  # CPU
                )
                self._initialized = True
                logger.info("Classification model loaded successfully")
            except Exception as e:
                logger.error(f"Failed to load classification model: {e}")
                # Fallback to a simpler model if distilbart fails
                try:
                    logger.info("Retrying with even smaller model...")
                    self._classifier = pipeline(
                        "zero-shot-classification",
                        model="typeform/distilbert-base-uncased-mnli",
                        device=-1
                    )
                    self._initialized = True
                except:
                    logger.error("All classification models failed to load")
                    raise
    
    # -----------------------------------------------------------------
    # Content extraction (handles all file types)
    # -----------------------------------------------------------------

    def _extract_content(self, file_path: str) -> Dict:
        """
        Extract content from *any* supported file type.

        Returns:
            dict with keys: full_text, page_count, paragraphs, headings
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        # --- PDF ---
        if ext == ".pdf":
            return self._extract_content_pdf(file_path)

        # --- ZIP (skip real analysis) ---
        if ext == ".zip":
            return {
                "full_text": "",
                "page_count": 0,
                "paragraphs": [],
                "headings": [],
            }

        # --- DOCX ---
        if ext == ".docx":
            return self._extract_content_docx(file_path)

        # --- Everything else: read as UTF-8 text ---
        return self._extract_content_text(file_path)

    def _extract_content_pdf(self, file_path: str) -> Dict:
        """Extract via pdf_processor and normalise into the common dict."""
        pdf_data = pdf_processor.extract_document(file_path)

        full_text = ""
        for page in pdf_data.pages:
            for block in page.blocks:
                for line in block.lines:
                    for span in line.spans:
                        full_text += span.text + " "

        paragraphs = [p.strip() for p in re.split(r"\n{2,}", full_text) if p.strip()]

        headings: List[str] = []
        for page in pdf_data.pages:
            for h in page.headings:
                headings.append(h.text)

        return {
            "full_text": full_text,
            "page_count": pdf_data.page_count,
            "paragraphs": paragraphs,
            "headings": headings,
            "_pdf_data": pdf_data,  # keep for legacy callers
        }

    def _extract_content_docx(self, file_path: str) -> Dict:
        """Extract text from DOCX, with fallback to raw read."""
        text = ""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            logger.warning("python-docx not installed, reading raw bytes as text")
            text = Path(file_path).read_text(encoding="utf-8", errors="replace")
        except Exception:
            text = Path(file_path).read_text(encoding="utf-8", errors="replace")

        paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
        headings = self._detect_headings_from_text(text)

        return {
            "full_text": text,
            "page_count": 1,
            "paragraphs": paragraphs,
            "headings": headings,
        }

    def _extract_content_text(self, file_path: str) -> Dict:
        """Read any text-based file as UTF-8."""
        try:
            text = Path(file_path).read_text(encoding="utf-8", errors="replace")
        except Exception:
            text = ""

        paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
        headings = self._detect_headings_from_text(text)

        return {
            "full_text": text,
            "page_count": 1,
            "paragraphs": paragraphs,
            "headings": headings,
        }

    @staticmethod
    def _detect_headings_from_text(text: str) -> List[str]:
        """Detect heading-like lines from plain text / markdown."""
        headings: List[str] = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            # Markdown headings
            if stripped.startswith("#"):
                headings.append(stripped.lstrip("# ").strip())
            # ALL-CAPS lines (at least 3 chars, not just symbols)
            elif len(stripped) >= 3 and stripped == stripped.upper() and re.search(r"[A-Z]", stripped):
                headings.append(stripped)
        return headings

    # -----------------------------------------------------------------
    # Content-signal analysis
    # -----------------------------------------------------------------

    def _compute_content_signals(self, text: str) -> Dict:
        """
        Analyze actual document text to compute structural signals used
        for content-signal-based chunking recommendations.

        Returns dict with keys:
            heading_density, code_ratio, table_ratio, list_ratio,
            avg_sentence_length, avg_paragraph_sentences,
            total_words, total_lines, total_paragraphs
        """
        lines = text.splitlines()
        total_lines = len(lines) or 1  # avoid division by zero

        # --- heading_density: lines starting with # or ALL CAPS ---------
        heading_lines = 0
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                heading_lines += 1
            elif len(stripped) >= 3 and stripped == stripped.upper() and re.search(r"[A-Z]", stripped):
                heading_lines += 1
        heading_density = heading_lines / total_lines

        # --- code_ratio: lines with code indicators ---------------------
        code_pattern = re.compile(
            r"(?:^\s*(?:def |class |function |import |from .+ import |const |let |var |"
            r"export |return |if |else |for |while |switch |case |try |catch |"
            r"async |await |=>|#!/))|```|[{}\[\]();]",
            re.IGNORECASE,
        )
        code_lines = sum(1 for l in lines if code_pattern.search(l))
        code_ratio = code_lines / total_lines

        # --- table_ratio: lines with 2+ pipe characters -----------------
        table_lines = sum(1 for l in lines if l.count("|") >= 2)
        table_ratio = table_lines / total_lines

        # --- list_ratio: lines starting with -, *, bullet, or numbered --
        list_pattern = re.compile(r"^\s*(?:[-*\u2022]|\d+[.)]\s)")
        list_lines = sum(1 for l in lines if list_pattern.match(l))
        list_ratio = list_lines / total_lines

        # --- sentence-level metrics -------------------------------------
        sentences = [s.strip() for s in re.split(r"[.!?]+", text) if s.strip()]
        sentence_lengths = [len(s.split()) for s in sentences]
        avg_sentence_length = float(np.mean(sentence_lengths)) if sentence_lengths else 0.0

        # --- paragraph-level metrics ------------------------------------
        paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
        total_paragraphs = len(paragraphs) or 1

        para_sentence_counts = []
        for para in paragraphs:
            para_sents = [s.strip() for s in re.split(r"[.!?]+", para) if s.strip()]
            para_sentence_counts.append(len(para_sents))
        avg_paragraph_sentences = float(np.mean(para_sentence_counts)) if para_sentence_counts else 0.0

        # --- totals -----------------------------------------------------
        words = text.split()
        total_words = len(words) or 1  # avoid division by zero

        # --- formula_ratio: lines with LaTeX or math operators ----------
        latex_pattern = re.compile(
            r"\$.*\$|\\frac|\\int|\\sum|\\begin\{equation\}",
        )
        math_operator_pattern = re.compile(
            r"(?:\d\s*[=+\-*/]\s*\d)"
        )
        formula_lines = 0
        for line in lines:
            if latex_pattern.search(line):
                formula_lines += 1
            elif math_operator_pattern.search(line):
                formula_lines += 1
        formula_ratio = formula_lines / total_lines

        # --- cross_ref_ratio: lines with cross-reference phrases --------
        cross_ref_pattern = re.compile(
            r"see\s+section|as\s+in\s+chapter|refer\s+to|as\s+mentioned\s+in|"
            r"as\s+shown\s+in|see\s+figure|see\s+table|cf\.|ibid|"
            r"subject\s+to\s+section|pursuant\s+to|in\s+section\s+\d|"
            r"under\s+section|per\s+section|obligations\s+in\s+section",
            re.IGNORECASE,
        )
        cross_ref_lines = sum(1 for l in lines if cross_ref_pattern.search(l))
        cross_ref_ratio = cross_ref_lines / total_lines

        # --- named_entity_density: consecutive capitalized words --------
        # Match 2+ consecutive capitalized words NOT at sentence start
        # We look for patterns mid-sentence (preceded by lowercase or punctuation + space)
        ne_pattern = re.compile(r"(?<=[a-z,;:]\s)(?:[A-Z][a-z]+\s+){1,}[A-Z][a-z]+")
        ne_matches = ne_pattern.findall(text)
        named_entity_count = sum(len(m.split()) for m in ne_matches)
        named_entity_density = named_entity_count / total_words

        # --- question_density: lines ending with ? ----------------------
        question_lines = sum(1 for l in lines if l.strip().endswith("?"))
        question_density = question_lines / total_lines

        # --- dialogue_ratio: lines with dialogue indicators -------------
        dialogue_pattern = re.compile(
            r'["\u201c\u201d]|'
            r"^\s*\w+\s*:|"
            r'\bQ\s*:|A\s*:|'
            r'\bsaid\b|\basked\b',
            re.IGNORECASE,
        )
        dialogue_lines = sum(1 for l in lines if dialogue_pattern.search(l))
        dialogue_ratio = dialogue_lines / total_lines

        # --- heading_depth: max depth of headings (1-6 scale) ----------
        heading_depth = 0
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                level = min(level, 6)
                heading_depth = max(heading_depth, level)
            elif (len(stripped) >= 3 and stripped == stripped.upper()
                  and re.search(r"[A-Z]", stripped) and heading_depth < 1):
                heading_depth = 1
        heading_depth = min(heading_depth, 6)

        # --- forward_references: lines with forward-looking phrases -----
        forward_ref_pattern = re.compile(
            r"we\s+will\s+see|later\s+in|in\s+the\s+next|upcoming\s+section|discussed\s+below",
            re.IGNORECASE,
        )
        forward_ref_lines = sum(1 for l in lines if forward_ref_pattern.search(l))
        forward_references = forward_ref_lines / total_lines

        # --- back_references: lines with backward-looking phrases -------
        back_ref_pattern = re.compile(
            r"as\s+mentioned|recall\s+that|as\s+we\s+saw|previously|as\s+discussed|in\s+the\s+previous",
            re.IGNORECASE,
        )
        back_ref_lines = sum(1 for l in lines if back_ref_pattern.search(l))
        back_references = back_ref_lines / total_lines

        # --- comparison_patterns: lines with comparison phrases ---------
        comparison_pattern = re.compile(
            r"\bvs\b|compared\s+to|unlike|in\s+contrast|whereas|on\s+the\s+other\s+hand|difference\s+between",
            re.IGNORECASE,
        )
        comparison_lines = sum(1 for l in lines if comparison_pattern.search(l))
        comparison_patterns = comparison_lines / total_lines

        # --- causal_chains: lines with causal phrases -------------------
        causal_pattern = re.compile(
            r"\bbecause\b|\btherefore\b|\bconsequently\b|as\s+a\s+result|leads\s+to|due\s+to|\bhence\b|\bthus\b",
            re.IGNORECASE,
        )
        causal_lines = sum(1 for l in lines if causal_pattern.search(l))
        causal_chains = causal_lines / total_lines

        # --- vocabulary_diversity: unique words / total words -----------
        lower_words = [w.lower() for w in words]
        unique_words = len(set(lower_words))
        vocabulary_diversity = unique_words / total_words

        return {
            "heading_density": round(heading_density, 4),
            "code_ratio": round(code_ratio, 4),
            "table_ratio": round(table_ratio, 4),
            "list_ratio": round(list_ratio, 4),
            "avg_sentence_length": round(avg_sentence_length, 1),
            "avg_paragraph_sentences": round(avg_paragraph_sentences, 1),
            "total_words": total_words,
            "total_lines": total_lines,
            "total_paragraphs": total_paragraphs,
            "formula_ratio": round(formula_ratio, 4),
            "cross_ref_ratio": round(cross_ref_ratio, 4),
            "named_entity_density": round(named_entity_density, 4),
            "question_density": round(question_density, 4),
            "dialogue_ratio": round(dialogue_ratio, 4),
            "heading_depth": heading_depth,
            "forward_references": round(forward_references, 4),
            "back_references": round(back_references, 4),
            "comparison_patterns": round(comparison_patterns, 4),
            "causal_chains": round(causal_chains, 4),
            "vocabulary_diversity": round(vocabulary_diversity, 4),
        }

    def _recommend_from_signals(self, signals: Dict, doc_type: str) -> Dict:
        """
        Priority-ordered decision tree that converts content signals into
        a concrete chunking recommendation.

        Returns dict with keys:
            chunking_method, chunk_size, overlap, embedding_model,
            reasoning, signals_used
        """
        embedding_model = "text-embedding-3-small"

        # Helper to append formula preservation note
        def _formula_note(reasoning: str) -> str:
            if signals.get("formula_ratio", 0) > 0.01:
                reasoning += (
                    f" Note: document contains mathematical formulas "
                    f"({signals['formula_ratio']:.1%} of lines). "
                    "Chunk boundaries should avoid splitting equations."
                )
            return reasoning

        # 1. Heavy code
        if signals.get("code_ratio", 0) > 0.3:
            return {
                "chunking_method": "code_aware",
                "chunk_size": 400,
                "overlap": 0,
                "embedding_model": embedding_model,
                "reasoning": _formula_note(
                    f"High code density ({signals['code_ratio']:.0%}). "
                    "Using code-aware chunking with zero overlap to preserve function boundaries."
                ),
                "signals_used": ["code_ratio"],
            }

        # 2. Deep hierarchical structure
        if signals.get("heading_depth", 0) >= 3 and signals.get("avg_paragraph_sentences", 0) > 3:
            return {
                "chunking_method": "hierarchical",
                "chunk_size": 700,
                "overlap": 100,
                "embedding_model": embedding_model,
                "reasoning": _formula_note(
                    f"Deep heading hierarchy (depth {signals['heading_depth']}) "
                    f"with substantial paragraphs ({signals['avg_paragraph_sentences']:.1f} sentences avg). "
                    "Hierarchical chunking preserves nested section structure."
                ),
                "signals_used": ["heading_depth", "avg_paragraph_sentences"],
            }

        # 3. Heading-rich prose
        if signals.get("heading_density", 0) > 0.03 and signals.get("avg_paragraph_sentences", 0) > 2:
            return {
                "chunking_method": "heading_based",
                "chunk_size": 600,
                "overlap": 75,
                "embedding_model": embedding_model,
                "reasoning": _formula_note(
                    f"Document has clear heading structure (density {signals['heading_density']:.1%}) "
                    f"with substantial paragraphs ({signals['avg_paragraph_sentences']:.1f} sentences avg). "
                    "Heading-based chunking preserves section boundaries."
                ),
                "signals_used": ["heading_density", "avg_paragraph_sentences"],
            }

        # 4. Question-heavy content (FAQ, interviews, Q&A)
        if signals.get("question_density", 0) > 0.05:
            return {
                "chunking_method": "sentence_window",
                "chunk_size": 300,
                "overlap": 40,
                "embedding_model": embedding_model,
                "reasoning": _formula_note(
                    f"High question density ({signals['question_density']:.1%} of lines). "
                    "Sentence-window chunking keeps each question-answer pair focused."
                ),
                "signals_used": ["question_density"],
            }

        # 5. Dense long-sentence prose (legal/academic)
        if signals.get("avg_sentence_length", 0) > 25:
            return {
                "chunking_method": "semantic",
                "chunk_size": 400,
                "overlap": 80,
                "embedding_model": embedding_model,
                "reasoning": _formula_note(
                    f"Dense prose with long sentences ({signals['avg_sentence_length']:.0f} words avg). "
                    "Semantic chunking with high overlap to avoid splitting complex clauses."
                ),
                "signals_used": ["avg_sentence_length"],
            }

        # 6. Short, fragmented text (chat logs, bullet lists, FAQs)
        if signals.get("avg_paragraph_sentences", 0) < 3 and signals.get("avg_sentence_length", 0) < 15:
            return {
                "chunking_method": "sentence_window",
                "chunk_size": 256,
                "overlap": 30,
                "embedding_model": embedding_model,
                "reasoning": _formula_note(
                    f"Short paragraphs ({signals['avg_paragraph_sentences']:.1f} sentences) "
                    f"and brief sentences ({signals['avg_sentence_length']:.0f} words). "
                    "Sentence-window chunking keeps each chunk focused."
                ),
                "signals_used": ["avg_paragraph_sentences", "avg_sentence_length"],
            }

        # 7. Table-heavy content
        if signals.get("table_ratio", 0) > 0.1:
            return {
                "chunking_method": "recursive",
                "chunk_size": 800,
                "overlap": 100,
                "embedding_model": embedding_model,
                "reasoning": _formula_note(
                    f"Significant table content ({signals['table_ratio']:.0%} of lines). "
                    "Larger recursive chunks to keep table rows together."
                ),
                "signals_used": ["table_ratio"],
            }

        # 8. Fallback
        return {
            "chunking_method": "recursive",
            "chunk_size": 512,
            "overlap": 50,
            "embedding_model": embedding_model,
            "reasoning": _formula_note(
                f"General {doc_type} document with no dominant structural signal. "
                "Using balanced recursive chunking."
            ),
            "signals_used": [],
        }

    # -----------------------------------------------------------------
    # Main analysis entry-point
    # -----------------------------------------------------------------

    async def analyze(self, document_path: str) -> Dict:
        """
        Analyze a document and recommend optimal RAG configuration.

        Args:
            document_path: Path to the document (any supported type)

        Returns:
            dict with keys: document_type, structure, density,
                          recommended_config, confidence_score, reasoning,
                          content_signals
        """
        import time
        start_time = time.time()
        logger.info(f"Analyzing document: {document_path}")

        ext = Path(document_path).suffix.lower()

        # --- ZIP shortcut: skip heavy analysis ---
        if ext == ".zip":
            return {
                "document_type": "general",
                "structure": {
                    "has_headings": False,
                    "has_tables": False,
                    "has_code_blocks": False,
                    "hierarchy_depth": 0,
                    "avg_paragraph_length": 0,
                },
                "density": {
                    "avg_sentence_length": 0,
                    "vocabulary_richness": 0,
                    "technical_term_density": 0,
                },
                "recommended_config": self._recommend_from_signals({}, "general"),
                "confidence_score": 0.5,
                "reasoning": "ZIP archive detected. Default configuration applied; extract individual files for better analysis.",
                "content_signals": {},
            }

        # Extract content (works for any file type)
        t0 = time.time()
        content = self._extract_content(document_path)
        logger.info(f"Content extraction complete in {time.time() - t0:.2f}s")

        full_text = content["full_text"]
        if not full_text.strip():
            logger.warning("Empty document text after extraction")

        # Get text sample for classification
        text_sample = full_text[:2000]

        # Use fast keyword-based classification first
        t1 = time.time()
        fast_result = self._quick_classify(text_sample)

        if fast_result:
            doc_type, confidence = fast_result
            logger.info(f"Quick classification: {doc_type} (confidence: {confidence}) in {time.time() - t1:.3f}s")
        else:
            # Fallback to ML-based classification
            logger.info("Keyword classification inconclusive. Falling back to ML...")
            doc_type, confidence = self._classify_document(text_sample)
            logger.info(f"ML classification: {doc_type} (confidence: {confidence}) in {time.time() - t1:.3f}s")

        logger.info(f"Classification took: {time.time() - t1:.2f}s")

        # Compute content signals from full text
        t2 = time.time()
        signals = self._compute_content_signals(full_text)
        logger.info(f"Content signal computation took: {time.time() - t2:.2f}s")

        # Signal-based recommendation
        t3 = time.time()
        config = self._recommend_from_signals(signals, doc_type)
        logger.info(f"Signal-based recommendation took: {time.time() - t3:.2f}s")

        # Build structure dict from signals
        structure = {
            "has_headings": signals["heading_density"] > 0,
            "has_tables": signals["table_ratio"] > 0,
            "has_code_blocks": signals["code_ratio"] > 0.05,
            "hierarchy_depth": 0,  # preserved for compat; heading depth not in signals
            "avg_paragraph_length": (
                int(signals["total_words"] / signals["total_paragraphs"])
                if signals["total_paragraphs"] else 0
            ),
        }

        # Enrich hierarchy_depth if content available
        lines = full_text.splitlines()
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                structure["hierarchy_depth"] = max(structure["hierarchy_depth"], level)

        # Build density dict from signals
        words = re.findall(r'\b\w+\b', text_sample.lower())
        vocabulary_richness = len(set(words)) / len(words) if words else 0
        long_words = [w for w in words if len(w) > 12]
        technical_term_density = len(long_words) / len(words) if words else 0

        density = {
            "avg_sentence_length": signals["avg_sentence_length"],
            "vocabulary_richness": round(vocabulary_richness, 2),
            "technical_term_density": round(technical_term_density, 2),
        }

        reasoning = config.pop("reasoning", "")
        signals_used = config.pop("signals_used", [])

        result = {
            "document_type": doc_type,
            "structure": structure,
            "density": density,
            "recommended_config": config,
            "confidence_score": confidence,
            "reasoning": reasoning,
            "content_signals": signals,
        }

        logger.info(f"Analysis complete in {time.time() - start_time:.2f}s: type={doc_type}, confidence={confidence:.2f}")
        return result
    
    def _quick_classify(self, text: str) -> Optional[tuple[str, float]]:
        """Fast keyword-based classification for obvious document types."""
        text_lower = text.lower()
        
        # Resume detection (very common for this app)
        resume_keywords = ["experience", "education", "skills", "projects", "certifications", "resume", "cv", "curriculum vitae"]
        # Check if text contains high density of resume sections
        matches = sum(1 for kw in resume_keywords if kw in text_lower)
        if matches >= 4:
            return "general", 0.95
            
        # Legal detection
        legal_keywords = ["hereby", "agreement", "notary", "jurisdiction", "confidentiality", "indemnification", "severability"]
        legal_matches = sum(1 for kw in legal_keywords if kw in text_lower)
        if legal_matches >= 3:
            return "legal", 0.9
            
        # Financial detection
        financial_keywords = ["balance sheet", "cash flow", "revenue", "fiscal year", "audit", "ebitda"]
        fin_matches = sum(1 for kw in financial_keywords if kw in text_lower)
        if fin_matches >= 2:
            return "financial", 0.85
            
        # Support/FAQ detection
        support_keywords = ["frequently asked questions", "how do i", "contact us", "support guide", "troubleshooting"]
        sup_matches = sum(1 for kw in support_keywords if kw in text_lower)
        if sup_matches >= 2:
            return "support", 0.85
        
        return None
    
    def _analyze_structure_from_text(self, content: Dict) -> Dict:
        """
        Analyze document structure from plain-text content dict
        (non-PDF path).
        """
        full_text = content.get("full_text", "")
        paragraphs = content.get("paragraphs", [])
        headings = content.get("headings", [])

        has_headings = len(headings) > 0

        # Detect tables: pipe-delimited rows or CSV-like structure
        lines = full_text.splitlines()
        pipe_lines = sum(1 for l in lines if l.count("|") >= 2)
        comma_lines = sum(1 for l in lines if l.count(",") >= 3)
        has_tables = pipe_lines >= 2 or comma_lines >= 5

        # Detect code
        code_indicators = ["```", "def ", "function ", "class ", "import ",
                           "const ", "var ", "let ", "=> {", "};", "#!/"]
        has_code = any(ind in full_text for ind in code_indicators)

        # Hierarchy depth from markdown heading levels
        hierarchy_depth = 0
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                hierarchy_depth = max(hierarchy_depth, level)

        # Average paragraph length
        avg_paragraph_length = (
            int(np.mean([len(p) for p in paragraphs])) if paragraphs else 0
        )

        return {
            "has_headings": has_headings,
            "has_tables": has_tables,
            "has_code_blocks": has_code,
            "hierarchy_depth": hierarchy_depth,
            "avg_paragraph_length": avg_paragraph_length,
        }
    
    def _classify_document(self, text_sample: str) -> tuple[str, float]:
        """
        Classify document type using zero-shot classification.
        
        Args:
            text_sample: Sample text from document
            
        Returns:
            tuple of (document_type, confidence_score)
        """
        self._ensure_initialized()
        
        categories = [
            "legal contract or agreement",
            "medical or healthcare document",
            "technical documentation or manual",
            "customer support or FAQ",
            "academic research paper",
            "financial report or statement",
            "general content or article"
        ]
        
        try:
            result = self._classifier(
                text_sample,
                candidate_labels=categories,
                multi_label=False
            )
            
            # Map to simplified labels
            label_map = {
                "legal contract or agreement": "legal",
                "medical or healthcare document": "medical",
                "technical documentation or manual": "technical",
                "customer support or FAQ": "support",
                "academic research paper": "academic",
                "financial report or statement": "financial",
                "general content or article": "general"
            }
            
            top_label = result["labels"][0]
            confidence = result["scores"][0]
            
            doc_type = label_map.get(top_label, "general")
            
            return doc_type, confidence
            
        except Exception as e:
            logger.warning(f"Classification failed, defaulting to 'general': {e}")
            return "general", 0.5
    
    def _analyze_structure(self, pdf_data) -> Dict:
        """
        Analyze document structure.
        
        Args:
            pdf_data: Extracted PDF data from pdf_processor
            
        Returns:
            dict with structure metrics
        """
        has_headings = False
        has_tables = False
        has_code_blocks = False
        hierarchy_depth = 0
        paragraph_lengths = []
        
        # Check for headings
        for page in pdf_data.pages:
            if page.headings:
                has_headings = True
                # Calculate hierarchy depth from heading levels
                levels = [h.level for h in page.headings]
                hierarchy_depth = max(hierarchy_depth, max(levels) if levels else 0)
            
            # Check for tables
            if page.tables:
                has_tables = True
            
            # Analyze paragraph lengths
            for block in page.blocks:
                block_text = ""
                for line in block.lines:
                    for span in line.spans:
                        block_text += span.text + " "
                
                if block_text.strip():
                    paragraph_lengths.append(len(block_text))
                    
                    # Detect code blocks (heuristic: monospace font, indentation)
                    if self._looks_like_code(block_text):
                        has_code_blocks = True
        
        avg_paragraph_length = (
            int(np.mean(paragraph_lengths)) if paragraph_lengths else 0
        )
        
        return {
            "has_headings": has_headings,
            "has_tables": has_tables,
            "has_code_blocks": has_code_blocks,
            "hierarchy_depth": hierarchy_depth,
            "avg_paragraph_length": avg_paragraph_length
        }
    
    def _looks_like_code(self, text: str) -> bool:
        """Heuristic to detect code blocks."""
        # Check for common code patterns
        code_patterns = [
            r'\bdef\s+\w+\(',  # Python functions
            r'\bfunction\s+\w+\(',  # JavaScript functions
            r'\bclass\s+\w+',  # Class definitions
            r'\bimport\s+\w+',  # Import statements
            r'[{}\[\]();]',  # Brackets and semicolons
            r'^\s{4,}',  # Heavy indentation
        ]
        
        for pattern in code_patterns:
            if re.search(pattern, text, re.MULTILINE):
                return True
        return False
    
    def _analyze_density(self, text: str) -> Dict:
        """
        Analyze text density and complexity.
        
        Args:
            text: Sample text from document
            
        Returns:
            dict with density metrics
        """
        # Split into sentences (simple approach)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        # Calculate average sentence length
        sentence_lengths = [len(s.split()) for s in sentences]
        avg_sentence_length = (
            np.mean(sentence_lengths) if sentence_lengths else 0
        )
        
        # Calculate vocabulary richness (unique words / total words)
        words = re.findall(r'\b\w+\b', text.lower())
        vocabulary_richness = (
            len(set(words)) / len(words) if words else 0
        )
        
        # Calculate technical term density (heuristic: words > 12 chars)
        long_words = [w for w in words if len(w) > 12]
        technical_term_density = (
            len(long_words) / len(words) if words else 0
        )
        
        return {
            "avg_sentence_length": round(avg_sentence_length, 1),
            "vocabulary_richness": round(vocabulary_richness, 2),
            "technical_term_density": round(technical_term_density, 2)
        }
    
    def _generate_config(
        self, 
        doc_type: str, 
        structure: Dict, 
        density: Dict
    ) -> Dict:
        """
        Generate recommended RAG configuration based on analysis.
        
        Args:
            doc_type: Classified document type
            structure: Structure analysis results
            density: Density analysis results
            
        Returns:
            dict with recommended configuration
        """
        # Base configurations by document type
        base_configs = {
            "legal": {
                "chunking_method": "semantic",
                "chunk_size": 900,
                "overlap": 125,
                "embedding_model": "text-embedding-3-small",
                "retrieval_k": 5,
                "reranking": {"provider": "cohere", "model": "rerank-english-v3.0", "return_k": 5, "top_n": 20}
            },
            "medical": {
                "chunking_method": "semantic",
                "chunk_size": 700,
                "overlap": 100,
                "embedding_model": "text-embedding-3-small",
                "retrieval_k": 5,
                "reranking": {"provider": "cohere", "model": "rerank-english-v3.0", "return_k": 5, "top_n": 20}
            },
            "technical": {
                "chunking_method": "semantic",
                "chunk_size": 600,
                "overlap": 80,
                "embedding_model": "text-embedding-3-small",
                "retrieval_k": 7,
                "reranking": {"provider": "cohere", "model": "rerank-english-v3.0", "return_k": 5, "top_n": 20}
            },
            "support": {
                "chunking_method": "character",
                "chunk_size": 400,
                "overlap": 60,
                "embedding_model": "text-embedding-3-small",
                "retrieval_k": 3,
                "reranking": None
            },
            "academic": {
                "chunking_method": "semantic",
                "chunk_size": 800,
                "overlap": 110,
                "embedding_model": "text-embedding-3-small",
                "retrieval_k": 6,
                "reranking": {"provider": "cohere", "model": "rerank-english-v3.0", "return_k": 5, "top_n": 20}
            },
            "financial": {
                "chunking_method": "semantic",
                "chunk_size": 700,
                "overlap": 100,
                "embedding_model": "text-embedding-3-small",
                "retrieval_k": 5,
                "reranking": {"provider": "cohere", "model": "rerank-english-v3.0", "return_k": 5, "top_n": 20}
            },
            "general": {
                "chunking_method": "character",
                "chunk_size": 600,
                "overlap": 80,
                "embedding_model": "text-embedding-3-small",
                "retrieval_k": 5,
                "reranking": None
            }
        }
        
        config = base_configs.get(doc_type, base_configs["general"]).copy()
        
        # Adjust based on structure
        if structure["has_tables"]:
            config["chunk_size"] = int(config["chunk_size"] * 1.2)
        
        if structure["has_code_blocks"]:
            config["chunking_method"] = "character"  # Better for code
        
        if structure["hierarchy_depth"] > 3:
            config["overlap"] = int(config["overlap"] * 1.3)
        
        # Adjust based on density
        if density["avg_sentence_length"] > 25:
            config["chunk_size"] = int(config["chunk_size"] * 1.15)
        
        if density["vocabulary_richness"] > 0.7:
            config["chunking_method"] = "semantic"
        
        if density["technical_term_density"] > 0.15:
            config["overlap"] = int(config["overlap"] * 1.2)
        
        return config
    
    def _explain_recommendation(
        self,
        doc_type: str,
        structure: Dict,
        density: Dict,
        config: Dict
    ) -> str:
        """
        Generate human-readable explanation for recommendations.
        
        Args:
            doc_type: Document type
            structure: Structure analysis
            density: Density analysis
            config: Recommended configuration
            
        Returns:
            Explanation string
        """
        explanations = []
        
        # Document type reasoning
        type_reasons = {
            "legal": "Legal documents require larger chunks to preserve clause context and contractual relationships.",
            "medical": "Medical documents need semantic chunking to maintain clinical context and HIPAA compliance.",
            "technical": "Technical documentation benefits from semantic chunking to keep related concepts together.",
            "support": "FAQ-style content works best with smaller, character-based chunks for exact matching.",
            "academic": "Academic papers need larger chunks to preserve citations and research context.",
            "financial": "Financial documents require semantic chunking to maintain numerical context and relationships.",
            "general": "General content uses balanced settings suitable for narrative text."
        }
        
        explanations.append(type_reasons.get(doc_type, type_reasons["general"]))
        
        # Structure adjustments
        if structure["has_tables"]:
            explanations.append("Chunk size increased by 20% to accommodate table structures.")
        
        if structure["has_code_blocks"] and config["chunking_method"] == "character":
            explanations.append("Using character-based chunking to preserve code formatting.")
        
        if structure["hierarchy_depth"] > 3:
            explanations.append(
                f"Increased overlap by 30% due to deep hierarchy (depth: {structure['hierarchy_depth']})."
            )
        
        # Density adjustments
        if density["avg_sentence_length"] > 25:
            explanations.append(
                f"Chunk size increased by 15% for complex sentences (avg: {density['avg_sentence_length']} words)."
            )
        
        if config["chunking_method"] == "semantic" and density["vocabulary_richness"] > 0.7:
            explanations.append(f"Semantic chunking selected due to high vocabulary richness ({density['vocabulary_richness']:.0%}).")
        
        if density["technical_term_density"] > 0.15:
            explanations.append(
                f"Increased overlap by 20% due to high technical term density ({density['technical_term_density']:.0%})."
            )
        
        return " ".join(explanations)


# Singleton instance
document_analyzer = DocumentAnalyzer()
